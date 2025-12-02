import glob
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown_it import MarkdownIt
from markdown_it.token import Token
from typing import List

# --- Configuration ---
# Use a specific style for code blocks
CODE_STYLE = 'Courier New' 
# --- End Configuration ---

def add_run_with_formatting(paragraph, text, bold=False, italic=False, code=False):
    """Adds a run to a paragraph with optional formatting."""
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if code:
        run.font.name = CODE_STYLE

def handle_inline(paragraph, tokens: List[Token]):
    """Handles inline markdown formatting like bold, italic, and code."""
    for token in tokens:
        if token.type == 'text':
            add_run_with_formatting(paragraph, token.content)
        elif token.type == 'strong_open':
            inline_tokens = token.children or []
            for t in inline_tokens:
                 add_run_with_formatting(paragraph, t.content, bold=True)
        elif token.type == 'em_open':
            inline_tokens = token.children or []
            for t in inline_tokens:
                 add_run_with_formatting(paragraph, t.content, italic=True)
        elif token.type == 'code_inline':
            add_run_with_formatting(paragraph, token.content, code=True)
        elif token.type == 'image':
            try:
                img_src = token.attrs.get('src')
                if img_src and os.path.exists(img_src):
                    paragraph.add_run().add_picture(img_src, width=Inches(4))
            except Exception as e:
                paragraph.add_run(f"[Image not found: {token.content}]")

def convert_md_to_docx(md_file_path: str):
    """
    Converts a single Markdown file to a DOCX document.
    """
    docx_file_path = os.path.splitext(md_file_path)[0] + '.docx'
    doc = Document()
    md = MarkdownIt()

    with open(md_file_path, 'r', encoding='utf-8') as f:
        md_content = f.read()

    tokens = md.parse(md_content)
    
    table = None
    list_level = 0
    
    for i, token in enumerate(tokens):
        if token.type == 'heading_open':
            level = int(token.tag[1])
            p = doc.add_heading('', level=level)
            inline_children = tokens[i+1].children or []
            handle_inline(p, inline_children)

        elif token.type == 'paragraph_open':
            p = doc.add_paragraph()
            inline_children = tokens[i+1].children or []
            handle_inline(p, inline_children)

        elif token.type == 'bullet_list_open':
            list_level += 1
        elif token.type == 'bullet_list_close':
            list_level -= 1
        
        elif token.type == 'ordered_list_open':
            list_level += 1
        elif token.type == 'ordered_list_close':
            list_level -= 1

        elif token.type == 'list_item_open':
            style = 'List Bullet' if list_level > 0 else 'List Number'
            p = doc.add_paragraph(style=style)
            inline_children = tokens[i+1].children or []
            handle_inline(p, inline_children)

        elif token.type == 'fence': # For code blocks like ```mermaid
            p = doc.add_paragraph(token.content.strip())
            p.style = 'Normal'
            p.paragraph_format.left_indent = Inches(0.5)
            font = p.runs[0].font
            font.name = CODE_STYLE
            font.size = Pt(10)

        elif token.type == 'table_open':
            # This is a basic table handler
            table_tokens = []
            j = i + 1
            while j < len(tokens) and tokens[j].type != 'table_close':
                table_tokens.append(tokens[j])
                j += 1
            
            header_row_tokens = table_tokens[2].children
            num_cols = len(header_row_tokens)
            
            # Create table
            table = doc.add_table(rows=1, cols=num_cols, style='Table Grid')
            hdr_cells = table.rows[0].cells
            for idx, th in enumerate(header_row_tokens):
                if th.type == 'th_open':
                    cell_content = table_tokens[2].children[idx+1].content
                    hdr_cells[idx].text = cell_content

            # Add body rows
            for row_token in table_tokens:
                if row_token.type == 'tr_open':
                    row_cells_tokens = row_token.children
                    # Skip header row
                    if row_cells_tokens and row_cells_tokens[0].tag == 'td':
                        row_cells = table.add_row().cells
                        for idx, td in enumerate(row_cells_tokens):
                             if td.type == 'td_open':
                                cell_content = row_cells_tokens[idx+1].content
                                row_cells[idx].text = cell_content
            # Skip past table tokens
            i = j
            table = None
            
    doc.save(docx_file_path)
    print(f"Successfully converted '{md_file_path}' to '{docx_file_path}'")

if __name__ == '__main__':
    md_files = glob.glob('*.md')
    for md_file in md_files:
        try:
            convert_md_to_docx(md_file)
        except Exception as e:
            print(f"Error converting {md_file}: {e}")
